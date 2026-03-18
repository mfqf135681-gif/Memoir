"""Word document importer."""

from pathlib import Path

from docx import Document

from .base import BaseImporter, register_importer


@register_importer
class WordImporter(BaseImporter):
    """Microsoft Word document importer."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def extract_text(self, file_path: str) -> str:
        """Extract text from a Word document.

        Args:
            file_path: Path to the Word file

        Returns:
            Extracted text content
        """
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)
