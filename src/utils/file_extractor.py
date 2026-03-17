"""File content extraction utilities for Memoir.

Supports TXT, PDF, DOCX, and images (via OCR).
"""

import io
from pathlib import Path
from typing import Optional

from ..utils.logger import get_logger

logger = get_logger(__name__)

# Optional imports
PDF_AVAILABLE = True
try:
    from pypdf import PdfReader
except ImportError:
    PDF_AVAILABLE = False

DOCX_AVAILABLE = True
try:
    from docx import Document
except ImportError:
    DOCX_AVAILABLE = False

PIL_AVAILABLE = True
try:
    from PIL import Image
except ImportError:
    PIL_AVAILABLE = False

TESSERACT_AVAILABLE = True
try:
    import pytesseract
except ImportError:
    TESSERACT_AVAILABLE = False


class FileExtractor:
    """Extract text content from various file formats."""

    @staticmethod
    async def extract(file_path: str, content_type: Optional[str] = None) -> str:
        """Extract text from a file.

        Args:
            file_path: Path to file
            content_type: Optional content type hint

        Returns:
            Extracted text content
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        # Determine content type if not provided
        if not content_type:
            content_type = FileExtractor._get_content_type(suffix)

        # Extract based on content type
        if content_type == "text/plain":
            return await FileExtractor._extract_text(path)
        elif content_type == "application/pdf":
            return await FileExtractor._extract_pdf(path)
        elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            return await FileExtractor._extract_docx(path)
        elif content_type.startswith("image/"):
            return await FileExtractor._extract_image(path)
        else:
            # Try to extract as text
            try:
                return await FileExtractor._extract_text(path)
            except Exception as e:
                logger.warning(f"Could not extract content from {path}: {e}")
                return ""

    @staticmethod
    def _get_content_type(suffix: str) -> str:
        """Get content type from file suffix.

        Args:
            suffix: File suffix (e.g., '.txt', '.pdf')

        Returns:
            Content type string
        """
        content_types = {
            ".txt": "text/plain",
            ".text": "text/plain",
            ".md": "text/markdown",
            ".markdown": "text/markdown",
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".webp": "image/webp",
        }
        return content_types.get(suffix, "application/octet-stream")

    @staticmethod
    async def _extract_text(file_path: Path) -> str:
        """Extract plain text.

        Args:
            file_path: Path to text file

        Returns:
            Text content
        """
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    async def _extract_pdf(file_path: Path) -> str:
        """Extract text from PDF.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF extraction not available. Install pypdf: pip install pypdf")

        text_parts = []
        reader = PdfReader(file_path)

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    @staticmethod
    async def _extract_docx(file_path: Path) -> str:
        """Extract text from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX extraction not available. Install python-docx: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @staticmethod
    async def _extract_image(file_path: Path) -> str:
        """Extract text from image using OCR.

        Args:
            file_path: Path to image file

        Returns:
            Extracted text
        """
        if not PIL_AVAILABLE:
            raise RuntimeError("Image processing not available. Install Pillow: pip install pillow")

        if not TESSERACT_AVAILABLE:
            raise RuntimeError("OCR not available. Install pytesseract: pip install pytesseract")

        # Open image
        image = Image.open(file_path)

        # Convert to RGB if needed
        if image.mode != "RGB":
            image = image.convert("RGB")

        # Perform OCR
        text = pytesseract.image_to_string(image)

        return text.strip()

    @staticmethod
    async def extract_from_bytes(
        content: bytes,
        filename: str,
        content_type: Optional[str] = None,
    ) -> str:
        """Extract text from bytes.

        Args:
            content: File content as bytes
            filename: Original filename
            content_type: Optional content type

        Returns:
            Extracted text
        """
        # Determine content type
        suffix = Path(filename).suffix.lower()
        if not content_type:
            content_type = FileExtractor._get_content_type(suffix)

        # Create temporary file
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            return await FileExtractor.extract(tmp_path, content_type)
        finally:
            # Clean up
            Path(tmp_path).unlink(missing_ok=True)


async def extract_file_content(
    content: bytes,
    filename: str,
    content_type: Optional[str] = None,
) -> str:
    """Convenience function to extract file content.

    Args:
        content: File content
        filename: Original filename
        content_type: Optional content type

    Returns:
        Extracted text
    """
    return await FileExtractor.extract_from_bytes(content, filename, content_type)
